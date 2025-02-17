import docx

docpath1 = 'word_documents/panda1.docx'
docpath2 = 'word_documents/panda2.docx'

doc1 = docx.Document(docpath1)
doc2 = docx.Document(docpath2)

text = doc2.paragraphs[0].text

new_paragraph = doc1.add_paragraph(text)

new_text = new_paragraph.text
print(new_text)

